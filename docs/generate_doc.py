#!/usr/bin/env python3
"""生成 VMS 项目模块文档 (DOCX)"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

# ============================================================
# 全局样式设置
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for level in range(1, 5):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Microsoft YaHei'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if level == 1:
        heading_style.font.size = Pt(22)
        heading_style.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    elif level == 2:
        heading_style.font.size = Pt(16)
        heading_style.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    elif level == 3:
        heading_style.font.size = Pt(13)
        heading_style.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

def add_code_block(doc, code_text):
    """添加代码块 (灰色背景段落)"""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(16)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x37, 0x37, 0x37)
        # 灰色背景
        shading_elm = p._element.get_or_add_pPr()
        shd = shading_elm.makeelement(qn('w:shd'), {
            qn('w:fill'): 'F5F5F5',
            qn('w:val'): 'clear',
        })
        shading_elm.append(shd)

def add_comment(doc, text):
    """添加注释行 (绿色, 缩进)"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"// {text}")
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    run.font.italic = True

def add_table_with_style(doc, headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Light Grid Accent 1')
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()  # spacing

# ============================================================
# 封面
# ============================================================
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('车辆管理系统 (VMS)\n模块功能与代码注释文档')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Vehicle Management System\n基于 Qt6 (C++17) 的车辆信息管理程序')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x61, 0x61, 0x61)

doc.add_paragraph()
doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('版本: v1.2\n').font.size = Pt(12)
info.add_run('作者: shadow\n').font.size = Pt(12)
info.add_run('仓库: https://github.com/shadow-bjut/VMS\n').font.size = Pt(12)
info.add_run(f'生成日期: 2026-07-13').font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 目录概览
# ============================================================
doc.add_heading('项目架构总览', level=1)
doc.add_paragraph('VMS 采用经典的三层架构设计，共分为 5 个模块：')

add_table_with_style(doc,
    ['模块', '文件', '层级', '职责'],
    [
        ['模块一: 程序入口', 'main.cpp', '启动层', 'QApplication 初始化、主窗口创建、事件循环'],
        ['模块二: 数据模型', 'vehicle.h/cpp, car.h/cpp, bus.h/cpp, truck.h/cpp', 'Model 层', 'Vehicle 抽象基类 + 3 个子类的继承体系'],
        ['模块三: 数据管理', 'vehiclemanager.h/cpp', 'Manager 层', 'CRUD 操作、搜索、统计、JSON 文件持久化'],
        ['模块四: 主窗口', 'mainwindow.h/cpp', 'View 层', '菜单/工具栏/表格/搜索面板/状态栏'],
        ['模块五: 对话框', 'addvehicledialog.h/cpp, editvehicledialog.h/cpp, statisticsdialog.h/cpp', 'View 层', '添加/编辑/统计三个子窗口'],
    ]
)

doc.add_paragraph('数据流方向：UI 对话框 → MainWindow → VehicleManager → Vehicle 派生类 → JSON 文件')
doc.add_page_break()

# ============================================================
# 模块一：程序入口
# ============================================================
doc.add_heading('模块一：程序入口 (main.cpp)', level=1)

doc.add_heading('主要功能', level=2)
doc.add_paragraph('程序的启动入口，负责 Qt 应用初始化、主窗口创建和事件循环管理。是整个系统的 "点火开关"。')

doc.add_heading('核心文件', level=2)
add_table_with_style(doc, ['文件', '行数', '说明'], [['main.cpp', '14', '程序入口，创建 QApplication 和 MainWindow']])

doc.add_heading('代码注释', level=2)

add_comment(doc, 'main.cpp — 程序入口')
add_code_block(doc, """#include <QApplication>
#include "mainwindow.h"

int main(int argc, char *argv[]) {
    // 1. 创建 QApplication 对象
    //    QApplication 是 Qt 应用程序的核心，管理事件循环和全局状态
    //    每个 Qt GUI 程序必须有且仅有一个 QApplication 实例
    QApplication app(argc, argv);

    // 2. 设置应用元数据
    //    这些信息会在 QMessageBox::about 等系统对话框中显示
    app.setApplicationName(QStringLiteral("VehicleManager"));
    app.setApplicationVersion(QStringLiteral("1.0"));
    app.setOrganizationName(QStringLiteral("SWE25"));

    // 3. 创建主窗口实例
    //    构造函数内部完成全部 UI 搭建 (setupUI)
    MainWindow w;

    // 4. 显示主窗口
    w.show();

    // 5. 进入 Qt 事件循环
    //    app.exec() 阻塞在此处，持续等待和分派用户交互事件
    //    当最后一个窗口关闭时，exec() 返回，程序退出
    return app.exec();
}""")

doc.add_paragraph()
doc.add_heading('设计要点', level=3)
doc.add_paragraph('• 整个 main.cpp 仅 14 行，体现了 Qt 框架的 "约定优于配置" 理念', style='List Bullet')
doc.add_paragraph('• QStringLiteral 宏确保中文字符串在编译期转换为 UTF-16，避免运行时编码转换开销', style='List Bullet')
doc.add_paragraph('• app.exec() 是同步阻塞调用，所有用户交互都在事件循环中以异步回调方式处理', style='List Bullet')
doc.add_page_break()

# ============================================================
# 模块二：数据模型层
# ============================================================
doc.add_heading('模块二：数据模型层 (Vehicle 继承体系)', level=1)

doc.add_heading('主要功能', level=2)
doc.add_paragraph('定义车辆数据结构和类型系统。采用抽象基类 + 多态的设计模式，向上层暴露统一接口 (Vehicle*)，隐藏具体车型的实现差异。')

doc.add_heading('核心文件', level=2)
add_table_with_style(doc,
    ['文件', '行数', '说明'],
    [
        ['vehicle.h', '82', 'Vehicle 抽象基类 + VehicleType 枚举 + 类型转换工具函数'],
        ['vehicle.cpp', '88', '基类实现：getter/setter、JSON 序列化/反序列化'],
        ['car.h', '26', 'Car 子类：小轿车，附加两厢/三厢属性'],
        ['car.cpp', '29', 'Car 实现：维护费 1000 元/月'],
        ['bus.h', '24', 'Bus 子类：大客车，附加载客量属性'],
        ['bus.cpp', '25', 'Bus 实现：维护费 2000 元/月'],
        ['truck.h', '24', 'Truck 子类：卡车，附加载重量属性'],
        ['truck.cpp', '25', 'Truck 实现：维护费 1500 元/月'],
    ]
)

doc.add_heading('类继承结构', level=3)

add_code_block(doc, """VehicleType (枚举)
  ├── Bus    (大客车)
  ├── Car    (小轿车)
  └── Truck  (卡车)

Vehicle (抽象基类)                    ← 9 个公共属性 + 2 个纯虚函数
  ├── Car   : public Vehicle          ← + m_trunkType (2=两厢/3=三厢)
  ├── Bus   : public Vehicle          ← + m_maxPassengers (载客量)
  └── Truck : public Vehicle          ← + m_maxLoad (吨)""")

doc.add_heading('代码注释', level=2)

add_comment(doc, 'vehicle.h — 抽象基类核心设计')
add_code_block(doc, """class Vehicle {
public:
    // --- 9 个公共属性的 getter/setter ---
    QString id() const;                    // 车辆编号 (唯一标识)
    void setId(const QString &id);
    QString plateNumber() const;           // 车牌号 (如 京A12345)
    QString manufacturer() const;          // 制造公司 (如一汽大众)
    QDate purchaseDate() const;            // 购买日期
    VehicleType vehicleType() const;       // 车辆类别 (只读, 构造时确定)
    double totalKm() const;                // 总公里数
    double fuelConsumption() const;        // 耗油量 (L/km)
    double basicMaintenance() const;       // 基本维护费 (元/月)
    double roadFee() const;                // 养路费
    double accumulatedCost() const;        // 累计总费用

    // 当月总费用 = 油价 × 耗油量 + 基本维护费
    // 纯虚函数 → 子类必须实现 (各车型可有不同计费公式)
    virtual double monthlyTotalCost(double fuelPrice) const = 0;

    // 获取车辆类型中文名称
    virtual QString typeName() const = 0;

    // JSON 序列化 / 反序列化 (虚函数, 支持子类扩展)
    virtual QJsonObject toJson() const;
    virtual void fromJson(const QJsonObject &obj);

protected:
    // 受保护构造函数 — 只允许子类调用
    // 子类通过此构造传入类型和默认维护费
    Vehicle(VehicleType type, double defaultMaintenance);
};""")

doc.add_paragraph()

add_comment(doc, 'car.h — 小轿车子类')
add_code_block(doc, """class Car : public Vehicle {
public:
    Car();  // 调用基类 Vehicle(VehicleType::Car, 1000.0)
            // 小轿车默认维护费 = 1000 元/月

    int trunkType() const;          // 获取厢数: 2=两厢, 3=三厢
    void setTrunkType(int t);       // 设置厢数 (只接受 2 或 3)
    QString trunkTypeName() const;  // 返回 "两厢" 或 "三厢"

    QString typeName() const override;                  // 返回 "小轿车"
    double monthlyTotalCost(double fuelPrice) const override;  // 油价×油耗 + 1000

    QJsonObject toJson() const override;    // 调用基类 → 追加 trunkType
    void fromJson(const QJsonObject &obj) override;  // 调用基类 → 读取 trunkType
private:
    int m_trunkType = 3; // 默认三厢
};""")

doc.add_paragraph()

add_comment(doc, 'bus.h — 大客车子类')
add_code_block(doc, """class Bus : public Vehicle {
public:
    Bus();  // 调用基类 Vehicle(VehicleType::Bus, 2000.0)
            // 大客车默认维护费 = 2000 元/月

    int maxPassengers() const;      // 最大载客量
    void setMaxPassengers(int n);

    QString typeName() const override;   // 返回 "大客车"
    double monthlyTotalCost(double fuelPrice) const override;
private:
    int m_maxPassengers = 0;
};""")

doc.add_paragraph()

add_comment(doc, 'truck.h — 卡车子类')
add_code_block(doc, """class Truck : public Vehicle {
public:
    Truck();  // 调用基类 Vehicle(VehicleType::Truck, 1500.0)
              // 卡车默认维护费 = 1500 元/月

    double maxLoad() const;         // 最大载重量 (吨)
    void setMaxLoad(double load);

    QString typeName() const override;   // 返回 "卡车"
    double monthlyTotalCost(double fuelPrice) const override;
private:
    double m_maxLoad = 0;
};""")

doc.add_paragraph()

add_comment(doc, 'vehicle.cpp — 枚举 ↔ 字符串转换 (支持中英文)')
add_code_block(doc, """// 枚举 → 中文字符串 (用于 JSON 序列化和表格显示)
QString vehicleTypeToString(VehicleType type) {
    switch (type) {
    case VehicleType::Bus:   return QStringLiteral("大客车");
    case VehicleType::Car:   return QStringLiteral("小轿车");
    case VehicleType::Truck: return QStringLiteral("卡车");
    }
    return QStringLiteral("未知");
}

// 字符串 → 枚举 (支持中文和英文输入, 用于 JSON 反序列化)
// fallback 为 VehicleType::Car (安全默认值)
VehicleType stringToVehicleType(const QString &str) {
    if (str == "大客车" || str == "Bus")   return VehicleType::Bus;
    if (str == "小轿车" || str == "Car")   return VehicleType::Car;
    if (str == "卡车"   || str == "Truck") return VehicleType::Truck;
    return VehicleType::Car;
}""")

doc.add_paragraph()

add_comment(doc, 'vehicle.cpp — JSON 序列化链 (模板方法模式)')
add_code_block(doc, """// 基类序列化：填充所有通用字段
QJsonObject Vehicle::toJson() const {
    QJsonObject obj;
    obj["id"]                = m_id;
    obj["plateNumber"]       = m_plateNumber;
    obj["manufacturer"]      = m_manufacturer;
    obj["purchaseDate"]      = m_purchaseDate.toString(Qt::ISODate);
    obj["vehicleType"]       = vehicleTypeToString(m_vehicleType);
    obj["totalKm"]           = m_totalKm;
    obj["fuelConsumption"]   = m_fuelConsumption;
    obj["basicMaintenance"]  = m_basicMaintenance;
    obj["roadFee"]           = m_roadFee;
    obj["accumulatedCost"]   = m_accumulatedCost;
    return obj;
}

// 子类序列化示例 (Car): 先调用基类，再追加自有字段
QJsonObject Car::toJson() const {
    QJsonObject obj = Vehicle::toJson();  // 第一步：获取基类字段
    obj["trunkType"] = m_trunkType;       // 第二步：追加子类特有字段
    return obj;
}""")

doc.add_heading('设计要点', level=3)
doc.add_paragraph('• 纯虚函数 monthlyTotalCost() 和 typeName() 强制子类实现，编译期保证多态完整性', style='List Bullet')
doc.add_paragraph('• toJson()/fromJson() 采用 "先基类后子类" 的调用链，实现代码复用（模板方法模式）', style='List Bullet')
doc.add_paragraph('• VehicleType 枚举使用 enum class 保证类型安全，避免隐式转换为 int', style='List Bullet')
doc.add_paragraph('• stringToVehicleType() 同时支持中英文输入，增强了 JSON 文件的兼容性', style='List Bullet')
doc.add_page_break()

# ============================================================
# 模块三：数据管理层
# ============================================================
doc.add_heading('模块三：数据管理层 (VehicleManager)', level=1)

doc.add_heading('主要功能', level=2)
doc.add_paragraph('系统的业务逻辑核心，管理 QVector<Vehicle*> 容器，提供 CRUD 操作、多种搜索策略、统计分析、JSON 文件持久化，以及油价和容量的全局配置管理。')

doc.add_heading('核心文件', level=2)
add_table_with_style(doc,
    ['文件', '行数', '说明'],
    [
        ['vehiclemanager.h', '80', 'VehicleManager 类声明 + Statistics 结构体'],
        ['vehiclemanager.cpp', '258', '全部业务逻辑实现'],
    ]
)

doc.add_heading('公有接口一览', level=3)
add_table_with_style(doc,
    ['分类', '方法', '返回值', '说明'],
    [
        ['增', 'addVehicle(Vehicle*, QString&)', 'bool', '添加车辆，4 级校验管道'],
        ['删', 'removeVehicle(id)', 'bool', '按编号删除 (线性搜索 + delete + removeAt)'],
        ['改', 'updateVehicle(oldId, newData, QString&)', 'bool', '更新车辆，条件式校验，JSON 深拷贝'],
        ['查', 'findById(id)', 'Vehicle*', '精确查找，返回单个指针或 nullptr'],
        ['查', 'findByPlateNumber(plate)', 'Vehicle*', '精确查找，返回单个指针或 nullptr'],
        ['查', 'findByPlateNumberFuzzy(keyword)', 'QVector<Vehicle*>', '模糊查找 (contains, 大小写不敏感)'],
        ['查', 'findByManufacturer(keyword)', 'QVector<Vehicle*>', '模糊查找 (contains, 大小写不敏感)'],
        ['查', 'findByType(type)', 'QVector<Vehicle*>', '按枚举类别筛选'],
        ['统计', 'getStatistics()', 'Statistics', '返回各类型计数'],
        ['统计', 'allVehicles()', 'const QVector<Vehicle*>&', '只读访问全部车辆'],
        ['文件', 'saveToFile(filePath, QString&)', 'bool', 'JSON 序列化写入文件 (新格式)'],
        ['文件', 'loadFromFile(filePath, QString&)', 'bool', 'JSON 反序列化加载 (兼容新旧格式)'],
        ['配置', 'setFuelPrice(price) / fuelPrice()', 'void / double', '油价管理 (影响当月费用计算)'],
        ['配置', 'setMaxCapacity(cap) / maxCapacity()', 'void / int', '容量管理 (下限=当前数量)'],
        ['工具', 'count() / isEmpty() / isFull()', 'int / bool / bool', '容量状态查询'],
        ['工具', 'clearAll()', 'void', '删除所有车辆并释放内存'],
    ]
)

doc.add_heading('代码注释', level=2)

add_comment(doc, 'vehiclemanager.cpp — addVehicle: 四级校验管道')
add_code_block(doc, """bool VehicleManager::addVehicle(Vehicle *vehicle, QString &errorMsg) {
    // 校验 1: 空指针检查
    if (!vehicle) {
        errorMsg = QStringLiteral("车辆对象为空！");
        return false;
    }
    // 校验 2: 容量上限检查
    if (isFull()) {
        errorMsg = QStringLiteral("车辆信息库已满，无法添加！");
        return false;
    }
    // 校验 3: 编号唯一性检查
    if (findById(vehicle->id())) {
        errorMsg = QStringLiteral("编号 ") + vehicle->id() + QStringLiteral(" 已存在，添加失败！");
        return false;
    }
    // 校验 4: 车牌号唯一性检查 (空车牌允许, 跳过检查)
    if (!vehicle->plateNumber().isEmpty() && findByPlateNumber(vehicle->plateNumber())) {
        errorMsg = QStringLiteral("车牌号 ") + vehicle->plateNumber() + QStringLiteral(" 已被使用，添加失败！");
        return false;
    }
    // 全部校验通过 → 加入容器
    m_vehicles.append(vehicle);
    return true;
}""")

doc.add_paragraph()

add_comment(doc, 'vehiclemanager.cpp — updateVehicle: 条件式校验 + JSON 深拷贝')
add_code_block(doc, """bool VehicleManager::updateVehicle(const QString &oldId, Vehicle *newData, QString &errorMsg) {
    if (!newData) {
        errorMsg = QStringLiteral("更新数据为空！");
        return false;
    }
    Vehicle *target = findById(oldId);
    if (!target) {
        errorMsg = QStringLiteral("编号 ") + oldId + QStringLiteral(" 不存在！");
        return false;
    }
    // 条件式校验：只有编号实际改变时才检查新编号冲突
    if (oldId != newData->id() && findById(newData->id())) {
        errorMsg = QStringLiteral("新编号 ") + newData->id() + QStringLiteral(" 已被使用！");
        return false;
    }
    // 条件式校验：只有车牌实际改变时才检查新车牌冲突
    if (target->plateNumber() != newData->plateNumber()
        && !newData->plateNumber().isEmpty()
        && findByPlateNumber(newData->plateNumber())) {
        errorMsg = QStringLiteral("新车牌号 ") + newData->plateNumber() + QStringLiteral(" 已被使用！");
        return false;
    }
    // JSON 深拷贝：保持 target 指针不变，避免外部持有的指针失效
    QJsonObject json = newData->toJson();
    target->fromJson(json);
    return true;
}""")

doc.add_paragraph()

add_comment(doc, 'vehiclemanager.cpp — loadFromFile: 格式兼容 + 工厂创建')
add_code_block(doc, """bool VehicleManager::loadFromFile(const QString &filePath, QString &errorMsg) {
    // ... 文件打开与 JSON 解析 (省略) ...

    clearAll(); // 清除现有数据，释放旧 Vehicle 对象内存

    // === 格式兼容判断 ===
    QJsonArray arr;
    if (doc.isArray()) {
        // 旧格式：顶层为纯数组，油价和容量使用默认值
        arr = doc.array();
    } else if (doc.isObject()) {
        // 新格式：顶层为对象，含 fuelPrice、maxCapacity、vehicles
        QJsonObject root = doc.object();
        m_fuelPrice = root["fuelPrice"].toDouble(m_fuelPrice);
        m_maxCapacity = root["maxCapacity"].toInt(m_maxCapacity);
        arr = root["vehicles"].toArray();
    } else {
        errorMsg = QStringLiteral("文件格式错误！");
        return false;
    }

    // 遍历 JSON 数组，根据 vehicleType 创建对应子类 (工厂模式)
    for (const QJsonValue &val : arr) {
        QJsonObject obj = val.toObject();
        VehicleType vt = stringToVehicleType(obj["vehicleType"].toString());
        Vehicle *v = nullptr;
        switch (vt) {
        case VehicleType::Bus:   v = new Bus();   break;
        case VehicleType::Car:   v = new Car();   break;
        case VehicleType::Truck: v = new Truck(); break;
        }
        if (v) {
            v->fromJson(obj);         // 多态反序列化
            m_vehicles.append(v);     // 加入容器
        }
    }
    return true;
}""")

doc.add_paragraph()

add_comment(doc, 'vehiclemanager.cpp — removeVehicle: 线性搜索 + 内存释放')
add_code_block(doc, """bool VehicleManager::removeVehicle(const QString &id) {
    for (int i = 0; i < m_vehicles.size(); ++i) {
        if (m_vehicles[i]->id() == id) {
            delete m_vehicles[i];       // 1. 释放堆内存 (VehicleManager 拥有所有权)
            m_vehicles.removeAt(i);     // 2. 从容器中移除指针
            return true;
        }
    }
    return false;
}""")

doc.add_heading('设计要点', level=3)
doc.add_paragraph('• VehicleManager 拥有 Vehicle 对象的所有权 (Owner of pointers)，负责 new/delete 生命周期管理', style='List Bullet')
doc.add_paragraph('• updateVehicle 采用 JSON 深拷贝而非替换指针，保证 MainWindow 中缓存的 Vehicle* 始终有效', style='List Bullet')
doc.add_paragraph('• loadFromFile 兼容旧格式 (纯数组) 和新格式 (含配置的对象)，实现向前兼容', style='List Bullet')
doc.add_paragraph('• setMaxCapacity 下限约束为当前车辆数，防止因容量缩小导致数据丢失', style='List Bullet')
doc.add_page_break()

# ============================================================
# 模块四：主窗口
# ============================================================
doc.add_heading('模块四：主窗口 (MainWindow)', level=1)

doc.add_heading('主要功能', level=2)
doc.add_paragraph('系统的交互中枢，负责搭建全部 UI 控件、连接信号槽、将用户操作转发给 VehicleManager，并管理窗口状态（标题栏文件路径、未保存标记）。')

doc.add_heading('核心文件', level=2)
add_table_with_style(doc,
    ['文件', '行数', '说明'],
    [
        ['mainwindow.h', '97', 'MainWindow 类声明 + 26 个槽函数 + UI 成员指针'],
        ['mainwindow.cpp', '~900', '全部 UI 搭建和交互逻辑'],
        ['mainwindow.ui', '—', 'Qt Designer 表单文件 (基础布局)'],
    ]
)

doc.add_heading('UI 组件树', level=3)
add_code_block(doc, """MainWindow (QMainWindow)
  ├── QMenuBar
  │     ├── 文件(&F): 打开 | 保存 | 另存为 | 退出
  │     ├── 操作(&V): 添加 | 编辑 | 删除 | 按编号删除 | 批量删除 | 油价 | 容量 | 显示全部
  │     ├── 统计(&T): 统计信息
  │     └── 帮助(&H): 关于 | 版本号
  ├── QToolBar (工具栏)
  │     └── 添加 | 编辑 | 删除 | 按编号删除 | 批量删除 | 打开 | 保存 | 另存为 | 统计
  ├── Central Widget
  │     ├── 搜索面板 (QGroupBox)
  │     │     ├── QLineEdit (搜索关键字)
  │     │     ├── QRadioButton × 4 (按编号/按车牌/按制造公司/按类别)
  │     │     ├── QPushButton "查询" (默认按钮, 回车触发)
  │     │     └── QPushButton "显示全部"
  │     └── QTableWidget (12列, 整行选择, 多选模式, 交替行颜色)
  └── QStatusBar
        └── QLabel × 5: 容量 | 总计 | 客车 | 轿车 | 卡车""")

doc.add_heading('代码注释', level=2)

add_comment(doc, 'mainwindow.cpp — 搜索流程: 4 种模式分发')
add_code_block(doc, """void MainWindow::onSearch() {
    QString keyword = m_searchEdit->text().trimmed();
    if (keyword.isEmpty()) {
        refreshTable();  // 空关键字 → 显示全部
        return;
    }
    QVector<Vehicle *> result;
    int mode = m_searchGroup->checkedId(); // 0=ID, 1=Plate, 2=Mfg, 3=Type

    switch (mode) {
    case 0: { // 按编号精确查找 → 最多返回1条
        Vehicle *v = m_manager.findById(keyword);
        if (v) result.append(v);
        else { /* 提示 "该编号不存在" */ return; }
        break;
    }
    case 1: { // 按车牌号模糊查找 (大小写不敏感)
        result = m_manager.findByPlateNumberFuzzy(keyword);
        if (result.isEmpty()) { /* 提示 "未找到匹配的车牌号" */ return; }
        break;
    }
    case 2: { // 按制造公司模糊查找 (大小写不敏感)
        result = m_manager.findByManufacturer(keyword);
        if (result.isEmpty()) { /* 提示 "该制造公司不存在" */ return; }
        break;
    }
    case 3: { // 按类别查找 → 先解析中文关键词, 再查询
        VehicleType vt;
        if (keyword.contains("客车") || keyword.contains("巴士"))
            vt = VehicleType::Bus;
        else if (keyword.contains("轿车"))
            vt = VehicleType::Car;
        else if (keyword.contains("卡车") || keyword.contains("货车"))
            vt = VehicleType::Truck;
        else { /* 提示 "请输入正确类别" */ return; }
        result = m_manager.findByType(vt);
        break;
    }
    }
    refreshTable(result); // 用搜索结果刷新表格
}""")

doc.add_paragraph()

add_comment(doc, 'mainwindow.cpp — 表格刷新: 填充 12 列 + 动态识别子类')
add_code_block(doc, """void MainWindow::refreshTable(const QVector<Vehicle *> &vehicles) {
    m_displayedVehicles = vehicles;
    m_table->setRowCount(vehicles.size());

    for (int row = 0; row < vehicles.size(); ++row) {
        Vehicle *v = vehicles[row];

        // 填充通用列 (0-9): id, plate, mfg, date, type, km, fuel, maintenance, road, acc
        setCell(0, v->id());
        // ... 中间 8 列 ...
        setCell(9, QString::number(v->accumulatedCost(), 'f', 2));

        // 列 10: 当月总费用 = 油价 × 油耗 + 维护费 (多态调用)
        double monthly = v->monthlyTotalCost(m_manager.fuelPrice());
        setCell(10, QString::number(monthly, 'f', 2));

        // 列 11: 特殊属性 (dynamic_cast 运行时类型识别)
        QString special;
        if (auto *bus = dynamic_cast<Bus *>(v)) {
            special = QStringLiteral("载客量: %1 人").arg(bus->maxPassengers());
        } else if (auto *car = dynamic_cast<Car *>(v)) {
            special = car->trunkTypeName();  // "两厢" 或 "三厢"
        } else if (auto *truck = dynamic_cast<Truck *>(v)) {
            special = QStringLiteral("载重量: %1 吨").arg(truck->maxLoad());
        }
        setCell(11, special);
    }
    updateStatusBar(); // 同步更新状态栏计数
}""")

doc.add_paragraph()

add_comment(doc, 'mainwindow.cpp — 排序: 12 列全部支持, lambda 比较器')
add_code_block(doc, """void MainWindow::onHeaderClicked(int section) {
    // 同一列再次点击 → 切换升降序；不同列 → 默认升序
    if (section == m_sortColumn)
        m_sortOrder = (m_sortOrder == Qt::AscendingOrder) ? Qt::DescendingOrder : Qt::AscendingOrder;
    else {
        m_sortColumn = section;
        m_sortOrder = Qt::AscendingOrder;
    }

    QVector<Vehicle *> sorted = m_displayedVehicles;
    bool asc = (m_sortOrder == Qt::AscendingOrder);

    switch (section) {
    case 0:  // 编号: 字符串比较
        std::sort(sorted.begin(), sorted.end(),
            [asc](Vehicle *a, Vehicle *b) { return asc ? a->id() < b->id() : a->id() > b->id(); });
        break;
    case 3:  // 购买时间: QDate 比较
        std::sort(sorted.begin(), sorted.end(),
            [asc](Vehicle *a, Vehicle *b) {
                return asc ? a->purchaseDate() < b->purchaseDate()
                           : a->purchaseDate() > b->purchaseDate(); });
        break;
    // ... 其他 10 列均有对应的比较器 ...
    case 10: // 当月总费用: 需要即时计算 (费用依赖于油价)
        std::sort(sorted.begin(), sorted.end(),
            [this, asc](Vehicle *a, Vehicle *b) {
                double ma = a->monthlyTotalCost(m_manager.fuelPrice());
                double mb = b->monthlyTotalCost(m_manager.fuelPrice());
                return asc ? ma < mb : ma > mb;
            });
        break;
    }
    refreshTable(sorted);
}""")

doc.add_paragraph()

add_comment(doc, 'mainwindow.cpp — 关闭事件: 未保存提醒')
add_code_block(doc, """void MainWindow::closeEvent(QCloseEvent *event) {
    if (!m_unsavedChanges) {
        event->accept();  // 无未保存修改 → 直接关闭
        return;
    }
    // 弹出三按钮对话框
    QMessageBox box(this);
    box.setWindowTitle(QStringLiteral("未保存的修改"));
    box.setText(QStringLiteral("当前有未保存的修改，是否保存后再退出？"));
    QPushButton *saveBtn    = box.addButton(QStringLiteral("保存"),   QMessageBox::AcceptRole);
    QPushButton *discardBtn = box.addButton(QStringLiteral("不保存"), QMessageBox::DestructiveRole);
    QPushButton *cancelBtn  = box.addButton(QStringLiteral("再想想"), QMessageBox::RejectRole);
    box.exec();

    QAbstractButton *clicked = box.clickedButton();
    if (clicked == saveBtn) {
        // 保存逻辑：有路径则直接保存，无路径则走另存为
        if (m_currentFilePath.isEmpty()) onSaveToFile();
        else onSave();
        if (m_unsavedChanges) { event->ignore(); return; } // 保存失败 → 不退出
        event->accept();
    } else if (clicked == discardBtn) {
        event->accept();  // 放弃修改 → 直接退出
    } else {
        event->ignore();  // 取消 → 返回事件循环
    }
}""")

doc.add_heading('设计要点', level=3)
doc.add_paragraph('• 使用 QButtonGroup 管理 4 个搜索模式的 RadioButton，checkedId() 一键获取当前模式', style='List Bullet')
doc.add_paragraph('• 搜索按类别时直接解析中文关键词 ("客车"/"巴士"/"轿车"/"卡车"/"货车")，比强制用户选枚举更友好', style='List Bullet')
doc.add_paragraph('• refreshTable 中 dynamic_cast 识别子类并显示特殊属性，体现了多态在 UI 层的应用', style='List Bullet')
doc.add_paragraph('• closeEvent 三按钮对话框 (保存/不保存/取消) 是 Qt 中处理未保存修改的标准模式', style='List Bullet')
doc.add_page_break()

# ============================================================
# 模块五：对话框层
# ============================================================
doc.add_heading('模块五：对话框层 (Dialogs)', level=1)

doc.add_heading('主要功能', level=2)
doc.add_paragraph('提供三个模态对话框，分别负责车辆数据的输入（添加）、修改（编辑）和可视化展示（统计）。采用表单验证 + 工厂方法 + QtCharts 图表。')

doc.add_heading('核心文件', level=2)
add_table_with_style(doc,
    ['文件', '行数', '说明'],
    [
        ['addvehicledialog.h', '69', 'AddVehicleDialog 类声明'],
        ['addvehicledialog.cpp', '242', '表单构建、校验、编号自动生成、createVehicle 工厂'],
        ['editvehicledialog.h', '55', 'EditVehicleDialog 类声明'],
        ['editvehicledialog.cpp', '231', '表单构建 + populateFields 预填 + createVehicle'],
        ['statisticsdialog.h', '20', 'StatisticsDialog 类声明'],
        ['statisticsdialog.cpp', '179', '统计计算 + 甜甜圈饼图 + 分组柱状图 (QtCharts)'],
    ]
)

doc.add_heading('代码注释', level=2)

add_comment(doc, 'addvehicledialog.cpp — 编号自动生成算法')
add_code_block(doc, """QString AddVehicleDialog::generateDefaultId(int typeIndex) const {
    // 1. 根据类型确定前缀
    QString prefix;
    switch (typeIndex) {
    case 0: prefix = QStringLiteral("BUS");   break;
    case 1: prefix = QStringLiteral("CAR");   break;
    case 2: prefix = QStringLiteral("TRUCK"); break;
    default: return QStringLiteral("XXX-001");
    }

    // 2. 扫描已有编号，找该前缀下的最大数字
    int maxNum = 0;
    for (const QString &id : m_existingIds) {
        if (id.startsWith(prefix + QStringLiteral("-"))) {
            QString numStr = id.mid(prefix.length() + 1); // 提取 "001" 部分
            bool ok = false;
            int num = numStr.toInt(&ok);
            if (ok && num > maxNum) maxNum = num;
        }
    }

    // 3. 生成下一个编号 (数字自增, 3位补零)
    //    例: 已有 BUS-001, BUS-003 → 生成 BUS-004
    return QStringLiteral("%1-%2").arg(prefix).arg(maxNum + 1, 3, 10, QChar('0'));
}""")

doc.add_paragraph()

add_comment(doc, 'addvehicledialog.cpp — 类别切换: 联动维护费 + 编号')
add_code_block(doc, """void AddVehicleDialog::onTypeChanged(int index) {
    // 1. 切换专属信息页面 (Bus → 载客量, Car → 厢数, Truck → 载重量)
    m_stacked->setCurrentIndex(index);

    // 2. 自动更新维护费默认值 (对应三种车型的月维护费)
    switch (index) {
    case 0: m_maintenanceSpin->setValue(2000.0); break; // 大客车: 2000 元/月
    case 1: m_maintenanceSpin->setValue(1000.0); break; // 小轿车: 1000 元/月
    case 2: m_maintenanceSpin->setValue(1500.0); break; // 卡车:   1500 元/月
    }

    // 3. 仅当用户未手动修改编号时，自动更新编号 (如 BUS-003 → CAR-001)
    if (m_idEdit->text().trimmed() == m_lastAutoId
        || m_idEdit->text().trimmed().isEmpty()) {
        m_lastAutoId = generateDefaultId(index);
        m_idEdit->setText(m_lastAutoId);
    }
}""")

doc.add_paragraph()

add_comment(doc, 'addvehicledialog.cpp — 输入校验 (lambda 闭包)')
add_code_block(doc, """// 确定按钮的校验闭包
connect(m_buttonBox, &QDialogButtonBox::accepted, this, [this]() {
    // 校验 1: 编号不能为空
    if (m_idEdit->text().trimmed().isEmpty()) {
        QMessageBox::warning(this, QStringLiteral("输入错误"),
                             QStringLiteral("车辆编号不能为空！"));
        return; // 阻止对话框关闭
    }

    // 校验 2: 车牌号长度必须为 7-8 个字符 (如 京A12345)
    QString plate = m_plateEdit->text().trimmed();
    if (!plate.isEmpty() && (plate.length() < 7 || plate.length() > 8)) {
        QMessageBox::warning(this, QStringLiteral("输入错误"),
                             QStringLiteral("车牌号必须合法！"));
        return; // 阻止对话框关闭
    }

    accept(); // 校验通过, 关闭对话框, 返回 QDialog::Accepted
});""")

doc.add_paragraph()

add_comment(doc, 'addvehicledialog.cpp — 工厂方法: 创建 Vehicle 子类对象')
add_code_block(doc, """Vehicle *AddVehicleDialog::createVehicle() const {
    int typeIdx = m_typeCombo->currentIndex();
    Vehicle *v = nullptr;

    // === 第一步: 根据类型创建对应子类, 并设置专属属性 ===
    switch (typeIdx) {
    case 0: { // 大客车
        auto *bus = new Bus();
        bus->setMaxPassengers(m_passengerSpin->value());  // 载客量
        v = bus;
        break;
    }
    case 1: { // 小轿车
        auto *car = new Car();
        car->setTrunkType(m_trunkCombo->currentData().toInt()); // 2=两厢, 3=三厢
        v = car;
        break;
    }
    case 2: { // 卡车
        auto *truck = new Truck();
        truck->setMaxLoad(m_loadSpin->value());  // 载重量(吨)
        v = truck;
        break;
    }
    default: return nullptr;
    }

    // === 第二步: 设置 9 个通用属性 ===
    v->setId(m_idEdit->text().trimmed());
    v->setPlateNumber(m_plateEdit->text().trimmed());
    v->setManufacturer(m_mfgEdit->text().trimmed());
    v->setPurchaseDate(m_purchaseEdit->date());
    v->setTotalKm(m_totalKmSpin->value());
    v->setFuelConsumption(m_fuelConsumptionSpin->value());
    v->setRoadFee(m_roadFeeSpin->value());
    v->setAccumulatedCost(m_accumulatedCostSpin->value());
    v->setBasicMaintenance(m_maintenanceSpin->value());

    return v; // 调用方拥有所有权, 负责 delete
}""")

doc.add_paragraph()

add_comment(doc, 'editvehicledialog.cpp — populateFields: 预填 + dynamic_cast')
add_code_block(doc, """void EditVehicleDialog::populateFields(const Vehicle *v) {
    if (!v) return;

    // === 第一步: 预填通用字段 (9 个 setter) ===
    m_idEdit->setText(v->id());
    m_plateEdit->setText(v->plateNumber());
    m_mfgEdit->setText(v->manufacturer());
    m_purchaseEdit->setDate(v->purchaseDate());
    m_totalKmSpin->setValue(v->totalKm());
    m_fuelConsumptionSpin->setValue(v->fuelConsumption());
    m_roadFeeSpin->setValue(v->roadFee());
    m_accumulatedCostSpin->setValue(v->accumulatedCost());
    m_maintenanceSpin->setValue(v->basicMaintenance());

    // === 第二步: 设置类型下拉框 ===
    int idx = 0;
    switch (v->vehicleType()) {
    case VehicleType::Bus:   idx = 0; break;
    case VehicleType::Car:   idx = 1; break;
    case VehicleType::Truck: idx = 2; break;
    }
    m_typeCombo->setCurrentIndex(idx);

    // === 第三步: 预填专属字段 (dynamic_cast 运行时类型识别) ===
    if (auto *bus = dynamic_cast<const Bus *>(v)) {
        m_passengerSpin->setValue(bus->maxPassengers());
    } else if (auto *car = dynamic_cast<const Car *>(v)) {
        int trunkIdx = (car->trunkType() == 2) ? 1 : 0;
        m_trunkCombo->setCurrentIndex(trunkIdx);
    } else if (auto *truck = dynamic_cast<const Truck *>(v)) {
        m_loadSpin->setValue(truck->maxLoad());
    }
}""")

doc.add_paragraph()

add_comment(doc, 'statisticsdialog.cpp — 甜甜圈饼图 + 分组柱状图 (QtCharts)')
add_code_block(doc, """void StatisticsDialog::setupUI(const Statistics &s,
                                double avgBusCost, double avgCarCost, double avgTruckCost) {
    // ===== 区域 1: 数量统计 (HTML QLabel) =====
    QGroupBox *numGroup = new QGroupBox(QStringLiteral("车辆数量统计"));
    // 4 个 QLabel, HTML 格式: <b>标题</b><br><span style='color:XX; font-size:24px'>计数</span> 辆
    // 总计(#333) | 大客车(#4CAF50) | 小轿车(#2196F3) | 卡车(#FF9800)

    if (s.total > 0) {
        // ===== 区域 2: 甜甜圈饼图 =====
        QPieSeries *pieSeries = new QPieSeries();
        pieSeries->setHoleSize(0.40);  // 40% 空洞 → 甜甜圈效果
        // 添加 3 个彩色切片, 设置标签 + 图例 + 动画 + 抗锯齿
        QChartView *pieView = new QChartView(pieChart);
        pieView->setRenderHint(QPainter::Antialiasing);

        // ===== 区域 3: 分组柱状图 (数量 + 平均月费用并排) =====
        QBarSet *setCount = new QBarSet(QStringLiteral("车辆数量"));
        setCount->setBrush(QColor("#42A5F5"));   // 蓝色柱
        *setCount << s.busCount << s.carCount << s.truckCount;

        QBarSet *setCost = new QBarSet(QStringLiteral("平均月费用"));
        setCost->setBrush(QColor("#FF7043"));    // 橙色柱
        *setCost << avgBusCost << avgCarCost << avgTruckCost;

        QBarSeries *barSeries = new QBarSeries();
        barSeries->append(setCount);
        barSeries->append(setCost);
        barSeries->setLabelsVisible(true);
        barSeries->setLabelsPosition(QAbstractBarSeries::LabelsOutsideEnd);

        // X 轴: 大客车 / 小轿车 / 卡车
        QBarCategoryAxis *axisX = new QBarCategoryAxis();
        axisX->append({"大客车", "小轿车", "卡车"});
        // Y 轴: 数量 / 元
        QValueAxis *axisY = new QValueAxis();
        axisY->setTitleText("数量 / 元");
    } else {
        // 无数据时显示空状态提示
        QLabel *emptyLabel = new QLabel("暂无车辆数据，请先添加车辆。");
    }

    // 确定按钮
    QDialogButtonBox *btn = new QDialogButtonBox(QDialogButtonBox::Ok);
}""")

doc.add_heading('设计要点', level=3)
doc.add_paragraph('• 编号自动生成采用 "按前缀分组 + 最大值 + 1" 策略，确保同类型车辆编号连续', style='List Bullet')
doc.add_paragraph('• 输入校验通过 lambda 闭包连接到 accepted 信号，校验失败时 return 阻止对话框关闭', style='List Bullet')
doc.add_paragraph('• AddVehicleDialog 和 EditVehicleDialog 的唯一差异在于编号处理：添加时自动生成，编辑时保留原值', style='List Bullet')
doc.add_paragraph('• StatisticsDialog 使用 QtCharts 模块 (QPieSeries + QBarSeries)，甜甜圈饼图 (holeSize=0.40) 比传统饼图更美观', style='List Bullet')
doc.add_paragraph('• 三种颜色方案 (绿/蓝/橙) 在统计图表和状态栏中保持一致，形成视觉识别体系', style='List Bullet')

doc.add_page_break()

# ============================================================
# 附录
# ============================================================
doc.add_heading('附录 A: 项目文件清单', level=1)
add_table_with_style(doc,
    ['分类', '文件', '行数', '说明'],
    [
        ['入口', 'main.cpp', '14', '程序启动入口'],
        ['模型', 'vehicle.h / vehicle.cpp', '82+88', 'Vehicle 抽象基类'],
        ['模型', 'car.h / car.cpp', '26+29', 'Car 小轿车'],
        ['模型', 'bus.h / bus.cpp', '24+25', 'Bus 大客车'],
        ['模型', 'truck.h / truck.cpp', '24+25', 'Truck 卡车'],
        ['管理', 'vehiclemanager.h / vehiclemanager.cpp', '80+258', 'VehicleManager 业务逻辑'],
        ['视图', 'mainwindow.h / mainwindow.cpp', '97+~900', 'MainWindow 主窗口'],
        ['视图', 'addvehicledialog.h / addvehicledialog.cpp', '69+242', '添加车辆对话框'],
        ['视图', 'editvehicledialog.h / editvehicledialog.cpp', '55+231', '编辑车辆对话框'],
        ['视图', 'statisticsdialog.h / statisticsdialog.cpp', '20+179', '统计对话框'],
        ['项目', 'VMS.pro', '38', 'Qt 项目文件 (widgets + charts)'],
    ]
)

doc.add_heading('附录 B: 流程图文件清单', level=1)
add_table_with_style(doc,
    ['模块', '流程图文件', '渲染命令'],
    [
        ['模块一: 程序入口', 'docs/entry-flow.dot', 'dot -Tpng docs/entry-flow.dot -o docs/entry-flow.png'],
        ['模块二: 数据模型', 'docs/vehicle-model-flow.dot', 'dot -Tpng docs/vehicle-model-flow.dot -o docs/vehicle-model-flow.png'],
        ['模块三: 数据管理', 'docs/vehiclemanager-flow.dot', 'dot -Tpng docs/vehiclemanager-flow.dot -o docs/vehiclemanager-flow.png'],
        ['模块四: 主窗口', 'docs/mainwindow-flow.dot', 'dot -Tpng docs/mainwindow-flow.dot -o docs/mainwindow-flow.png'],
        ['模块五: 对话框', 'docs/dialogs-flow.dot', 'dot -Tpng docs/dialogs-flow.dot -o docs/dialogs-flow.png'],
    ]
)

doc.add_heading('附录 C: 技术栈', level=1)
add_table_with_style(doc,
    ['技术', '版本/说明'],
    [
        ['语言', 'C++17'],
        ['框架', 'Qt 6.11.1 (widgets + charts)'],
        ['编译器', 'MinGW 64-bit'],
        ['构建系统', 'qmake (VMS.pro)'],
        ['序列化', 'QJsonDocument / QJsonObject / QJsonArray'],
        ['图表', 'QtCharts (QPieSeries, QBarSeries)'],
        ['流程图', 'Graphviz DOT (dot -Tpng)'],
    ]
)

# ============================================================
# 保存
# ============================================================
output_path = os.path.join(os.path.dirname(__file__), 'VMS_模块功能与代码注释文档.docx')
doc.save(output_path)
print(f"文档已保存到: {output_path}")
