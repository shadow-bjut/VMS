#!/usr/bin/env python3
"""生成 VMS 项目模块文档 — 精简版 (仅功能描述, 不含代码)"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ============================================================
# 全局样式
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if level == 1:
        hs.font.size = Pt(22)
        hs.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    elif level == 2:
        hs.font.size = Pt(16)
        hs.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    elif level == 3:
        hs.font.size = Pt(13)
        hs.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Light Grid Accent 1')
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

# ============================================================
# 封面
# ============================================================
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('车辆管理系统 (VMS)\n模块功能说明文档')
r.font.size = Pt(28)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

doc.add_paragraph()
s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('Vehicle Management System\n基于 Qt6 + C++17 的车辆信息管理程序')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x61, 0x61, 0x61)

doc.add_paragraph()
doc.add_paragraph()
i = doc.add_paragraph()
i.alignment = WD_ALIGN_PARAGRAPH.CENTER
i.add_run('版本: v1.2  |  作者: shadow  |  日期: 2026-07-13').font.size = Pt(11)

doc.add_page_break()

# ============================================================
# 架构总览
# ============================================================
doc.add_heading('项目架构总览', 1)
doc.add_paragraph('VMS 采用三层架构，将系统划分为 5 个功能模块。数据流方向为：UI 对话框 → MainWindow → VehicleManager → Vehicle 派生类 → JSON 文件。')

add_table(doc,
    ['模块', '文件', '层级', '职责'],
    [
        ['模块一: 程序入口', 'main.cpp', '启动层', 'QApplication 初始化、主窗口创建、事件循环管理'],
        ['模块二: 数据模型', 'vehicle/car/bus/truck (.h/.cpp)', 'Model 层', 'Vehicle 抽象基类 + 3 子类的继承与多态体系'],
        ['模块三: 数据管理', 'vehiclemanager (.h/.cpp)', 'Manager 层', '车辆容器的 CRUD、搜索、统计、JSON 文件持久化、全局配置'],
        ['模块四: 主窗口', 'mainwindow (.h/.cpp)', 'View 层', '菜单栏/工具栏/搜索面板/车辆表格/状态栏的搭建与交互调度'],
        ['模块五: 对话框', 'add/edit/statistics dialog (.h/.cpp)', 'View 层', '添加/编辑车辆表单、统计图表展示'],
    ]
)

doc.add_page_break()

# ============================================================
# 模块一
# ============================================================
doc.add_heading('模块一：程序入口', 1)

doc.add_heading('功能概述', 2)
doc.add_paragraph('main.cpp 是程序的唯一启动入口，仅 14 行代码。负责创建 Qt 应用实例、设置应用元数据（名称"VehicleManager"、版本"1.0"、组织名"SWE25"）、实例化主窗口并进入事件循环。app.exec() 为阻塞调用，程序在此等待并分派所有用户交互事件，直到最后一个窗口关闭后退出。')

doc.add_heading('执行流程', 2)
bullet('创建 QApplication 对象 — Qt GUI 程序的核心，管理全局事件循环')
bullet('设置应用元数据 — 应用名称、版本号、组织名（显示于关于对话框）')
bullet('创建 MainWindow 实例 — 构造函数内部完成全部 UI 搭建')
bullet('调用 w.show() 显示主窗口')
bullet('进入 app.exec() 事件循环 — 阻塞等待用户交互，窗口关闭后返回并退出')

doc.add_heading('关键设计', 2)
bullet('使用 QStringLiteral 宏确保中文字符串在编译期转换为 UTF-16，避免运行时编码转换开销')
bullet('整个入口仅 14 行，体现了 Qt 框架"约定优于配置"的设计理念')
bullet('app.exec() 为同步阻塞，所有用户操作以异步信号槽方式在事件循环中处理')

doc.add_page_break()

# ============================================================
# 模块二
# ============================================================
doc.add_heading('模块二：数据模型层 (Vehicle 继承体系)', 1)

doc.add_heading('功能概述', 2)
doc.add_paragraph('定义车辆数据结构和类型系统。采用抽象基类 Vehicle + 3 个派生类（Car/Bus/Truck）的继承体系。基类封装 9 个公共属性及对应的 getter/setter，并声明 2 个纯虚函数供子类实现。子类在基类基础上追加类型专属属性。整个继承体系通过 Vehicle* 基类指针向上层暴露统一接口，实现运行时多态。')

doc.add_heading('类结构', 2)

add_table(doc,
    ['类', '月维护费', '专属属性', 'typeName 返回值'],
    [
        ['Vehicle (抽象基类)', '—', '9 个公共属性: id, plateNumber, manufacturer, purchaseDate, totalKm, fuelConsumption, basicMaintenance, roadFee, accumulatedCost', '纯虚函数'],
        ['Car : Vehicle', '1000 元/月', 'm_trunkType: 2=两厢, 3=三厢 (默认三厢)', '"小轿车"'],
        ['Bus : Vehicle', '2000 元/月', 'm_maxPassengers: 最大载客量 (默认0)', '"大客车"'],
        ['Truck : Vehicle', '1500 元/月', 'm_maxLoad: 最大载重量, 单位吨 (默认0)', '"卡车"'],
    ]
)

doc.add_heading('VehicleType 枚举', 2)
doc.add_paragraph('定义 Bus / Car / Truck 三个枚举值，配套两个工具函数：vehicleTypeToString() 将枚举转为中文字符串（序列化时使用），stringToVehicleType() 将中/英文字符串转为枚举（反序列化时使用，同时兼容"Bus"/"Car"/"Truck"英文输入，fallback 为 Car）。')

doc.add_heading('核心虚函数', 2)

bullet('monthlyTotalCost(double fuelPrice) → double — 计算当月总费用 = 油价 × 耗油量 + 基本维护费。三个子类当前使用相同公式，但设计为虚函数允许未来各车型差异化计费')
bullet('typeName() → QString — 返回车辆类型中文名称（"大客车"/"小轿车"/"卡车"），用于表格和统计显示')
bullet('toJson() → QJsonObject — 序列化为 JSON。子类调用链：先调用 Vehicle::toJson() 获取基类 9 字段，再追加自有字段（模板方法模式）')
bullet('fromJson(const QJsonObject&) — 从 JSON 反序列化。调用链与 toJson() 对称：先调基类读通用字段，再读子类专属字段')

doc.add_heading('关键设计', 2)
bullet('纯虚函数强制子类实现，编译期保证多态完整性')
bullet('enum class VehicleType 保证类型安全，避免隐式 int 转换')
bullet('toJson/fromJson 采用"先基类后子类"调用链，最大化代码复用')
bullet('stringToVehicleType 同时支持中英文输入，增强 JSON 文件兼容性')

doc.add_page_break()

# ============================================================
# 模块三
# ============================================================
doc.add_heading('模块三：数据管理层 (VehicleManager)', 1)

doc.add_heading('功能概述', 2)
doc.add_paragraph('系统的业务逻辑核心，管理 QVector<Vehicle*> 容器（默认容量 100 辆）。提供完整的 CRUD 操作、5 种搜索策略、分类统计、JSON 文件持久化，以及油价和容量的全局配置管理。VehicleManager 拥有容器内所有 Vehicle 对象的所有权，负责 new/delete 生命周期管理。')

doc.add_heading('公有接口', 2)

add_table(doc,
    ['分类', '方法', '返回', '说明'],
    [
        ['增', 'addVehicle(Vehicle*, QString&)', 'bool', '4 级校验: 空指针→容量满→编号重复→车牌重复, 全部通过后 append'],
        ['删', 'removeVehicle(id)', 'bool', '线性搜索匹配 id, delete + removeAt, 未找到返回 false'],
        ['改', 'updateVehicle(oldId, newData, QString&)', 'bool', '条件式校验(仅检查实际修改的字段) + JSON 中转深拷贝'],
        ['查', 'findById(id)', 'Vehicle*', '精确匹配, 返回单个指针或 nullptr'],
        ['查', 'findByPlateNumber(plate)', 'Vehicle*', '精确匹配, 返回单个指针或 nullptr'],
        ['查', 'findByPlateNumberFuzzy(kw)', 'QVector<Vehicle*>', 'contains 模糊匹配, 大小写不敏感'],
        ['查', 'findByManufacturer(kw)', 'QVector<Vehicle*>', 'contains 模糊匹配, 大小写不敏感'],
        ['查', 'findByType(VehicleType)', 'QVector<Vehicle*>', '按枚举类型筛选'],
        ['统计', 'getStatistics()', 'Statistics', '遍历统计 total/busCount/carCount/truckCount'],
        ['统计', 'allVehicles()', 'const QVector<Vehicle*>&', '只读访问全部车辆'],
        ['文件', 'saveToFile(path, QString&)', 'bool', '遍历序列化为 JSON 对象 {fuelPrice, maxCapacity, vehicles[]}'],
        ['文件', 'loadFromFile(path, QString&)', 'bool', '解析 JSON, 兼容旧格式(纯数组)和新格式(对象), 工厂创建子类'],
        ['配置', 'setFuelPrice(p) / fuelPrice()', 'void/double', '油价管理, 默认 7.5 元/升, 影响当月费用计算'],
        ['配置', 'setMaxCapacity(n) / maxCapacity()', 'void/int', '容量管理, 下限=当前车辆数, 默认 100'],
        ['工具', 'clearAll()', 'void', 'delete 所有 Vehicle 并清空容器'],
    ]
)

doc.add_heading('addVehicle 校验流程', 2)
bullet('校验 1: vehicle == nullptr → 返回错误 "车辆对象为空"')
bullet('校验 2: isFull() (当前数量 >= 最大容量) → 返回错误 "信息库已满"')
bullet('校验 3: findById() 非空 (编号已存在) → 返回错误 "编号重复"')
bullet('校验 4: 车牌号非空且 findByPlateNumber() 非空 → 返回错误 "车牌已被使用"')
bullet('全部通过 → m_vehicles.append(vehicle), 返回 true')

doc.add_heading('updateVehicle 深拷贝策略', 2)
doc.add_paragraph('与 addVehicle 不同，updateVehicle 采用条件式校验（只有编号/车牌实际改变时才检查冲突），更新方式为 JSON 中转深拷贝（newData→toJson→fromJson→target），而非替换指针。这样 MainWindow 中缓存的 Vehicle* 指针始终有效，不会成为悬空指针。')

doc.add_heading('loadFromFile 格式兼容', 2)
doc.add_paragraph('加载时先判断 JSON 顶层类型：若为数组则按旧格式处理（纯车辆列表，油价/容量使用默认值），若为对象则按新格式处理（含 fuelPrice、maxCapacity、vehicles 三个字段）。遍历 vehicles 数组时，根据 vehicleType 字段通过 switch 创建对应的 Bus/Car/Truck 子类实例（工厂模式），再调用多态 fromJson 填充数据。')

doc.add_heading('关键设计', 2)
bullet('VehicleManager 拥有 Vehicle 对象所有权，析构函数自动调用 clearAll() 释放全部内存')
bullet('updateVehicle 使用 JSON 深拷贝而非替换指针，保护外部持有的引用')
bullet('loadFromFile 兼容新旧两种 JSON 格式，实现向前兼容')
bullet('setMaxCapacity 下限约束为当前车辆数，防止容量缩小导致数据丢失')
bullet('所有写操作通过 errorMsg 输出参数返回详细错误信息')

doc.add_page_break()

# ============================================================
# 模块四
# ============================================================
doc.add_heading('模块四：主窗口 (MainWindow)', 1)

doc.add_heading('功能概述', 2)
doc.add_paragraph('系统的交互中枢，约 900 行代码。负责搭建全部 UI 控件（菜单栏、工具栏、搜索面板、12 列表格、状态栏），连接 20+ 个信号槽，将用户操作转发给 VehicleManager 处理，并管理窗口状态（标题栏显示当前文件路径、m_unsavedChanges 标记追踪未保存修改）。')

doc.add_heading('UI 组件结构', 2)

add_table(doc,
    ['区域', '组件', '内容'],
    [
        ['菜单栏', '4 个 QMenu', '文件(打开/保存/另存为/退出) | 操作(添加/编辑/删除/按编号删除/批量删除/油价/容量/显示全部) | 统计(统计信息) | 帮助(关于/版本号)'],
        ['工具栏', 'QToolBar', '10 个快捷按钮: 添加/编辑/删除/按编号删除/批量删除 | 打开/保存/另存为 | 统计'],
        ['搜索面板', 'QGroupBox', 'QLineEdit + 4×QRadioButton(按编号/车牌/制造公司/类别) + 查询按钮 + 显示全部按钮。回车键触发查询。'],
        ['车辆表格', 'QTableWidget', '12 列(编号/车牌/制造公司/购买时间/类别/总公里数/耗油量/维护费/养路费/累计费用/当月费用/特殊属性)。整行选择+多选模式+交替行颜色+隐藏行号。'],
        ['状态栏', 'QStatusBar', '5 个 QLabel: 容量(x/100) | 总计(x辆) | 客车(x辆) | 轿车(x辆) | 卡车(x辆)'],
    ]
)

doc.add_heading('核心交互流程', 2)

doc.add_heading('搜索流程 (onSearch)', 3)
doc.add_paragraph('读取 QButtonGroup 的 checkedId() 确定搜索模式（0=按编号/1=按车牌/2=按制造公司/3=按类别），空关键字直接显示全部。模式 0~2 调用 VehicleManager 对应查找方法，模式 3 先解析中文关键词（支持"客车"/"巴士"/"轿车"/"卡车"/"货车"）再按类型查询。结果为空时弹出提示，有结果时调用 refreshTable() 展示。')

doc.add_heading('表格刷新 (refreshTable)', 3)
doc.add_paragraph('参数为 QVector<Vehicle*> 结果集。遍历时填充 12 列：前 10 列为通用属性（通过 Vehicle* 基类指针直接访问），第 10 列调用虚函数 monthlyTotalCost(fuelPrice) 计算当月费用，第 11 列通过 dynamic_cast 运行时识别子类类型并显示专属属性（载客量/两厢三厢/载重量）。填充完毕后调用 updateStatusBar() 同步状态栏计数。')

doc.add_heading('排序 (onHeaderClicked)', 3)
doc.add_paragraph('点击表头触发。同一列再次点击切换升降序，不同列默认升序。通过 setSortIndicator() 更新表头箭头方向。12 列各有一个 std::sort lambda 比较器，按对应字段类型排序（字符串/日期/数值）。排序在当前显示结果集上进行，排序后调用 refreshTable() 更新表格。')

doc.add_heading('删除操作 (三种方式)', 3)

add_table(doc,
    ['方式', '触发途径', '前序检查', '确认内容'],
    [
        ['选中行删除', '菜单/工具栏/右键', '信息库空? → 选中行?', '简单确认对话框, 显示编号'],
        ['按编号删除', '菜单/工具栏', '信息库空? → 输入编号 → 编号存在?', '详细确认对话框, 显示车牌/类别/制造商'],
        ['批量删除', '菜单/工具栏/Delete键', '是否选中行?', '列出前10条ID + "不可撤销"警告'],
    ]
)

doc.add_heading('文件操作流程', 3)
bullet('保存 (onSave): 若无文件路径则走另存为流程；有路径则直接写入。成功后将 m_unsavedChanges 置为 false，更新窗口标题')
bullet('另存为 (onSaveToFile): QFileDialog 选择路径后写入，成功后记录 m_currentFilePath')
bullet('打开 (onLoadFromFile): QFileDialog 选择文件后调用 VehicleManager::loadFromFile，成功后 refreshTable() 并更新窗口标题')

doc.add_heading('关闭保护 (closeEvent)', 3)
doc.add_paragraph('重写 closeEvent：若 m_unsavedChanges 为 false 则直接接受关闭。否则弹出三按钮对话框（保存 / 不保存 / 取消）。保存按钮根据是否有文件路径决定直接保存或另存为，保存失败则阻止退出（event->ignore()）。不保存直接退出。取消回到事件循环。')

doc.add_heading('右键上下文菜单', 3)
doc.add_paragraph('在表格上右键弹出菜单。始终包含"添加车辆"。若点击在有效行上，追加"编辑""删除""复制编号""复制车牌号"。复制功能通过 QApplication::clipboard() 写入系统剪贴板。若点击行不在选中范围内，自动选中该行。')

doc.add_heading('关键设计', 2)
bullet('QButtonGroup 管理 4 个 RadioButton，checkedId() 一键获取搜索模式')
bullet('按类别搜索支持中文关键词自然输入，比强制选枚举更友好')
bullet('refreshTable 中 dynamic_cast 识别子类展示专属信息，多态在 UI 层的典型应用')
bullet('closeEvent 三按钮对话框是 Qt 中处理未保存修改的标准实践')
bullet('m_unsavedChanges 标记贯穿所有写操作，确保数据安全')

doc.add_page_break()

# ============================================================
# 模块五
# ============================================================
doc.add_heading('模块五：对话框层 (Dialogs)', 1)

doc.add_heading('功能概述', 2)
doc.add_paragraph('提供三个模态对话框，分别负责车辆数据的输入（AddVehicleDialog）、修改（EditVehicleDialog）和可视化展示（StatisticsDialog）。两个编辑对话框采用相同的表单结构（通用信息区 + QStackedWidget 专属信息区），差异仅在于初始化和类别切换行为。统计对话框使用 QtCharts 模块展示甜甜圈饼图和分组柱状图。')

doc.add_heading('对话框对比', 2)

add_table(doc,
    ['特性', 'AddVehicleDialog', 'EditVehicleDialog', 'StatisticsDialog'],
    [
        ['用途', '录入新车辆', '修改已有车辆', '查看统计数据'],
        ['构造函数参数', 'existingIds (QStringList)', 'const Vehicle* (只读)', 'const VehicleManager& (只读)'],
        ['初始化方式', '自动生成默认编号', 'populateFields() 预填数据', '查询统计+计算平均费用'],
        ['类别切换行为', '更新维护费 AND 自动更新编号', '更新维护费 (不 touch 编号)', '—'],
        ['校验规则', '编号非空 + 车牌 7~8 字符', '同添加 (规则一致)', '—'],
        ['产出', 'Vehicle* (调用方拥有所有权)', 'Vehicle* (调用方拥有所有权)', '无 (只读展示)'],
        ['界面组件', '表单: QLineEdit + SpinBox + ComboBox + QStackedWidget', '同添加 (结构完全一致)', 'QtCharts: QPieSeries + QBarSeries'],
    ]
)

doc.add_heading('编号自动生成算法 (AddVehicleDialog)', 3)
doc.add_paragraph('generateDefaultId(typeIndex) 根据车辆类型确定前缀（BUS/CAR/TRUCK），遍历已有编号集合 m_existingIds，筛选同前缀编号并提取数字部分找最大值，生成"前缀-(max+1)"格式的新编号（3位补零）。例如已有 BUS-001 和 BUS-003，则生成 BUS-004。')

doc.add_heading('类别切换联动 (AddVehicleDialog)', 3)
doc.add_paragraph('onTypeChanged(index) 执行三个联动操作：① 切换 QStackedWidget 到对应专属信息页面（载客量/厢数/载重量）；② 自动更新维护费默认值为对应车型的标准值（2000/1000/1500）；③ 判断当前编号是否为自动生成或为空，若是则重新生成新类型的默认编号，若用户已手动修改则保留。')

doc.add_heading('输入校验机制', 3)
doc.add_paragraph('两个编辑对话框共享相同的校验逻辑。确定按钮的 accepted 信号连接到一个 lambda 闭包：先检查编号是否为空字符串（trimmed 后），再检查车牌号长度是否合法（非空时必须在 7~8 字符范围）。任一校验失败则弹出 QMessageBox::warning 并 return（阻止 accept() 调用，对话框保持打开）。全部通过后才调用 accept() 关闭对话框。')

doc.add_heading('populateFields 预填 (EditVehicleDialog)', 3)
doc.add_paragraph('接收 const Vehicle* 参数（只读、不修改原数据），依次填充通用字段、设置类型下拉框索引、通过 dynamic_cast 识别子类并填充专属字段。使用 const 语义确保编辑对话框不会意外修改正在被编辑的原始对象。')

doc.add_heading('createVehicle 工厂方法', 3)
doc.add_paragraph('两个编辑对话框共享相同的 createVehicle() 实现。两步流程：第一步根据 ComboBox 索引 switch 创建对应子类（Bus/Car/Truck）并设置专属属性；第二步统一调用 9 个 setter 填充通用属性。返回的 Vehicle* 指针所有权转移给调用方，调用方负责 delete。')

doc.add_heading('统计图表 (StatisticsDialog)', 3)
doc.add_paragraph('构造函数从 VehicleManager 获取统计数据并计算各类型平均月费用。setupUI 分三个区域：')
bullet('区域 1 — 数量统计: 4 个 HTML 格式 QLabel，以彩色大字显示总计/大客车(绿)/小轿车(蓝)/卡车(橙)的计数')
bullet('区域 2 — 甜甜圈饼图: QPieSeries 设置 holeSize=0.40 产生甜甜圈效果，三种颜色切片对应三类车辆，白色标签+bottom 图例+动画+抗锯齿')
bullet('区域 3 — 分组柱状图: QBarSeries 包含两组数据（蓝色柱=车辆数量，橙色柱=平均月费用），X 轴为三类车型，Y 轴标签"数量/元"，支持标签外侧显示数值')
bullet('无数据时显示灰色居中提示"暂无车辆数据，请先添加车辆"')

doc.add_heading('关键设计', 2)
bullet('AddVehicleDialog 和 EditVehicleDialog 共享表单结构，减少 UI 代码重复')
bullet('编号自动生成采用"前缀分组+最大值+1"策略，确保同类型编号有序连续')
bullet('校验失败通过 return 阻止 accept()，对话框保持打开等待用户修正')
bullet('EditVehicleDialog 使用 const Vehicle* 确保预填时不修改原数据')
bullet('统计图表三种颜色(绿/蓝/橙)与状态栏保持一致，建立统一视觉语言')
bullet('甜甜圈饼图 holeSize=0.40 比传统饼图更现代美观')

doc.add_page_break()

# ============================================================
# 附录
# ============================================================
doc.add_heading('附录 A: 项目文件清单', 1)

add_table(doc,
    ['文件', '行数', '所属模块', '说明'],
    [
        ['main.cpp', '14', '模块一', '程序入口'],
        ['vehicle.h', '82', '模块二', 'Vehicle 抽象基类 + VehicleType 枚举'],
        ['vehicle.cpp', '88', '模块二', '基类实现 + 枚举转换函数'],
        ['car.h / car.cpp', '26+29', '模块二', 'Car 小轿车 (维护费 1000)'],
        ['bus.h / bus.cpp', '24+25', '模块二', 'Bus 大客车 (维护费 2000)'],
        ['truck.h / truck.cpp', '24+25', '模块二', 'Truck 卡车 (维护费 1500)'],
        ['vehiclemanager.h / vehiclemanager.cpp', '80+258', '模块三', '数据管理业务逻辑'],
        ['mainwindow.h / mainwindow.cpp', '97+898', '模块四', '主窗口 UI 与交互'],
        ['addvehicledialog.h / addvehicledialog.cpp', '69+242', '模块五', '添加车辆对话框'],
        ['editvehicledialog.h / editvehicledialog.cpp', '55+231', '模块五', '编辑车辆对话框'],
        ['statisticsdialog.h / statisticsdialog.cpp', '20+179', '模块五', '统计对话框'],
        ['VMS.pro', '38', '项目配置', 'Qt qmake 项目文件'],
    ]
)

doc.add_heading('附录 B: 流程图索引', 1)

add_table(doc,
    ['模块', 'DOT 文件', 'PNG 渲染命令'],
    [
        ['模块一', 'docs/entry-flow.dot', 'dot -Tpng docs/entry-flow.dot -o docs/entry-flow.png'],
        ['模块二', 'docs/vehicle-model-flow.dot', 'dot -Tpng docs/vehicle-model-flow.dot -o docs/vehicle-model-flow.png'],
        ['模块三', 'docs/vehiclemanager-flow.dot', 'dot -Tpng docs/vehiclemanager-flow.dot -o docs/vehiclemanager-flow.png'],
        ['模块四', 'docs/mainwindow-flow.dot', 'dot -Tpng docs/mainwindow-flow.dot -o docs/mainwindow-flow.png'],
        ['模块五', 'docs/dialogs-flow.dot', 'dot -Tpng docs/dialogs-flow.dot -o docs/dialogs-flow.png'],
    ]
)

doc.add_heading('附录 C: 技术栈', 1)

add_table(doc,
    ['技术', '版本/说明'],
    [
        ['编程语言', 'C++17'],
        ['UI 框架', 'Qt 6.11.1 (widgets 模块 + charts 模块)'],
        ['编译器', 'MinGW 64-bit (GCC)'],
        ['构建系统', 'qmake (VMS.pro)'],
        ['数据序列化', 'QJsonDocument / QJsonObject / QJsonArray'],
        ['图表库', 'QtCharts (QPieSeries, QBarSeries, QChartView)'],
        ['流程图工具', 'Graphviz (dot -Tpng)'],
        ['文档格式', 'python-docx 生成 .docx'],
        ['版本控制', 'Git (https://github.com/shadow-bjut/VMS)'],
    ]
)

# ============================================================
# 保存
# ============================================================
out = os.path.join(os.path.dirname(__file__), 'VMS_模块功能说明_精简版.docx')
doc.save(out)
print(f"文档已保存到: {out}")
